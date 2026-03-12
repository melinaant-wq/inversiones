"""
Rebuild completo del documento de Inversiones.
Todo escrito con la API nativa de python-docx — sin copiar XML.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = '/Users/melinaant/Downloads/Definiciones Inversiones (1).docx'

# ── helpers ──────────────────────────────────────────────────────────────────

def h(doc, text, level=1):
    return doc.add_heading(text, level=level)

def p(doc, text='', bold=False, italic=False):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    return para

def tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = hd
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864'); tcPr.append(shd)
        for para in c.paragraphs:
            for r in para.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    if widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths): cell.width = Inches(widths[i])
    return t

def mono(doc, text):
    stns = [s.name for s in doc.styles if s.name]
    para = doc.add_paragraph()
    para.style = doc.styles['No Spacing' if 'No Spacing' in stns else 'Normal']
    r = para.add_run(text); r.font.name = 'Courier New'; r.font.size = Pt(9)
    pPr = para._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2'); pPr.append(shd)
    return para

def note(doc, text):
    para = doc.add_paragraph(); r = para.add_run(text); r.italic = True
    pPr = para._p.get_or_add_pPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB'); pPr.append(shd)
    return para

def sep(doc): doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin = Inches(1); s.bottom_margin = Inches(1)
    s.left_margin = Inches(1.25); s.right_margin = Inches(1.25)

# ══ PORTADA ══════════════════════════════════════════════════════════════════
t0 = doc.add_heading('Inversiones — Definición de Producto', 0)
t0.alignment = WD_ALIGN_PARAGRAPH.CENTER
p(doc, 'Este documento describe el módulo de Inversiones de Dinero Mágico: qué hace, cómo está estructurado, y qué entra en V1 vs V2. Está pensado para que alguien que no vio los prototipos pueda entender el producto completo.', italic=True)
sep(doc)

# ══ QUÉ ES ═══════════════════════════════════════════════════════════════════
h(doc, '¿Qué es el módulo de Inversiones?')
p(doc, 'El módulo de Inversiones le permite al usuario comprar y vender fracciones de acciones y ETFs estadounidenses desde la app de Dinero Mágico, usando pesos argentinos o dólares. El usuario puede ver su portafolio en tiempo real, explorar el mercado, analizar activos individuales y seguir el estado de sus órdenes — todo dentro de la misma app.')
sep(doc)
p(doc, 'El módulo está compuesto por 9 secciones:', bold=True)
tbl(doc, ['#', 'Nombre', 'Descripción'], [
    ('M0', 'Onboarding',         'Primera experiencia: perfil de inversor, creación de cuenta y recomendación inicial.'),
    ('M1', 'Home / Portafolio',  'Dashboard principal: balance, distribución, posiciones y gráfico.'),
    ('M2', 'Mercado',            'Pantalla de descubrimiento: Packs, ETFs y Acciones disponibles.'),
    ('M3', 'Detalle del activo', 'Vista individual: precio, gráfico, info, posición propia.'),
    ('M4', 'Flujo de compra',    'Proceso para ejecutar una orden de compra.'),
    ('M5', 'Flujo de venta',     'Proceso para ejecutar una orden de venta.'),
    ('M6', 'Centro de órdenes',  'Registro de todas las órdenes y sus estados.'),
    ('M7', 'Resultados',         'Historial de posiciones cerradas con G/P realizado.'),
    ('M8', 'Reportes y Comunicaciones', 'Reportes descargables (mensual/anual) y resúmenes periódicos de posición.'),
], widths=[0.6, 1.8, 4.2])
sep(doc)

# ══ MAPA ═════════════════════════════════════════════════════════════════════
h(doc, 'Mapa de navegación')
p(doc, 'Cómo se conectan los módulos entre sí y cuáles son los puntos de entrada principales.')
sep(doc)
mono(doc, """\
PRIMERA VEZ EN INVERSIONES
        │
        ▼
┌──────────────────────────────────────────────────────────┐
│             ONBOARDING (M0)                              │
│  Propuesta de valor → Cuestionario → Resultado → Acción  │
└──────────────────────────┬───────────────────────────────┘
         (sesiones futuras: acceso directo)
                           │
                           ▼
┌──────────────────────────────────────────────────────────┐
│              INVERSIONES HOME (M1)                       │
│     Balance · Distribución · Posiciones · Gráfico        │
└──────────┬───────────────────────────┬───────────────────┘
           │                           │
           ▼                           ▼
 ┌──────────────────┐       ┌────────────────────────┐
 │ CENTRO DE        │       │    RESULTADOS (M7)     │
 │ ÓRDENES (M6)     │       │  Posiciones cerradas   │
 └──────────────────┘       └────────────────────────┘

┌──────────────────────────────────────────────────────────┐
│                    MERCADO (M2)                          │
│         Packs · ETFs · Acciones · Buscador               │
└──────────────────────────┬───────────────────────────────┘
                           ▼
┌──────────────────────────────────────────────────────────┐
│              DETALLE DEL ACTIVO (M3)                     │
│    Gráfico de precio · Info · Posición · Historial       │
└───────────────────┬──────────────────────┬───────────────┘
                    │                      │
                    ▼                      ▼
      ┌─────────────────────┐  ┌──────────────────────┐
      │ FLUJO DE COMPRA (M4)│  │  FLUJO DE VENTA (M5) │
      └──────────┬──────────┘  └──────────┬───────────┘
                 └──────────┬─────────────┘
                            ▼
             ┌──────────────────────────────┐
             │    CENTRO DE ÓRDENES (M6)    │
             └──────────────────────────────┘

  M7 (Resultados) y M8 (Reportes) accesibles desde el header de M1.""")
sep(doc)

# ══ RESUMEN V1 vs V2 ═════════════════════════════════════════════════════════
h(doc, 'Resumen V1 vs V2')
p(doc, '✅  V1 — Lo que sale al lanzamiento', bold=True)
tbl(doc, ['Módulo', 'Qué incluye en V1'], [
    ('Onboarding (M0)',         'Cuestionario de perfil (5 preguntas), resultado con Pack recomendado, modo exploración sin operar, persistencia del progreso, reasignación desde Configuración.'),
    ('Home / Portafolio (M1)', 'Balance con selector de períodos, barra de distribución, lista de posiciones, indicador de orden activa, vista gráfico con comparación.'),
    ('Mercado (M2)',            'Packs, ETFs, Acciones con filtros y buscador en tiempo real.'),
    ('Detalle del activo (M3)','Gráfico de precio, descripción, noticias (Alpaca/ES), posición del usuario, historial de transacciones, favoritos.'),
    ('Flujo de compra (M4)',    'Orden de mercado en ARS o USD, resumen con comisión, confirmación.'),
    ('Flujo de venta (M5)',     'Orden de mercado en ARS o USD, resumen con comisión, confirmación.'),
    ('Centro de órdenes (M6)', 'Órdenes de mercado: Procesando → Completa / Cancelada.'),
    ('Resultados (M7)',         'Posiciones cerradas con G/P realizado, rendimiento total, detalle expandible.'),
    ('Reportes y Comunicaciones (M8)', 'Reporte mensual y anual en PDF. Comunicación mensual automática (push + email) con P&L realizado y no realizado.'),
], widths=[2.2, 4.4])
sep(doc)
p(doc, '🔜  V2 — Próximas iteraciones', bold=True)
tbl(doc, ['Funcionalidad', 'Descripción'], [
    ('Órdenes límite (M6)',            'Precio objetivo por el usuario. Estados: Activa → Ejecutada / Vencida / Cancelada.'),
    ('After market (M6)',              'Operar fuera del horario regular. Ventana: 4:00pm–8:00pm ET.'),
    ('Re-evaluación de perfil (M0)',   'Recordatorio periódico para actualizar el perfil de inversor.'),
    ('Perfil granular (M0)',           '5 perfiles en lugar de 3.'),
    ('Marcador precio promedio (M3)',  'Línea del precio promedio del usuario sobre el gráfico del activo.'),
    ('Valoración / Fundamentals (M3)','P/E, Market Cap, márgenes, flujo de caja. Requiere proveedor externo.'),
    ('Noticias en tiempo real (M3)',   'Reemplaza Alpaca por feed financiero real. Requiere proveedor externo.'),
    ('Reporte en Excel/CSV (M8)',      'Descarga alternativa al PDF para análisis en planilla.'),
    ('Reporte con formato AFIP (M8)', 'Estructura para declaración impositiva.'),
    ('Comunicación configurable (M8)','Frecuencia (semanal/mensual/trimestral) y canal configurables por el usuario.'),
    ('Comparación vs benchmark (M8)', 'Rendimiento vs S&P 500 en resúmenes y reportes.'),
], widths=[2.4, 4.2])
sep(doc)
p(doc, '⏳  Dependencias externas', bold=True)
tbl(doc, ['Proveedor', 'Qué provee', 'Cuándo'], [
    ('Alpaca',                 'Ejecución de órdenes, precios en tiempo real, noticias (inglés)',          'V1'),
    ('Proveedor FX',           'Tipo de cambio ARS/USD',                                                   'V1'),
    ('Proveedor fundamentals', 'Valoración, estados financieros, noticias en tiempo real',                 'Post-lanzamiento'),
], widths=[1.8, 3.8, 1.0])
sep(doc)

# ══ FLUJO E2E ════════════════════════════════════════════════════════════════
h(doc, 'Flujo completo de una compra (V1)')
sep(doc)
mono(doc, """\
Usuario en Home de Inversiones
        │
        ▼
Toca "Comprar" o va a Mercado (M2)
        │
        ▼
Selecciona activo → Detalle del activo (M3)
  · Ve precio actual y gráfico histórico · Toca "Comprar"
        │
        ▼
Flujo de compra (M4)
  1. Selecciona moneda: ARS o USD
  2. Ingresa monto (si supera saldo → error "Saldo insuficiente")
  3. Revisa: cantidad estimada · monto · comisión
  4. Confirmar compra
        │
        ▼
Centro de Órdenes (M6) — Estado inicial: Procesando 🟡
        ├── Éxito → Completa 🟢 (precio real + acciones recibidas)
        └── Error  → Cancelada 🔴 (motivo + Reintentar)
        │
        ▼
Banner en Home actualizado · Posición aparece en portafolio""")
sep(doc)

# ══ ESPECIFICACIÓN ═══════════════════════════════════════════════════════════
doc.add_page_break()
h(doc, 'Especificación detallada por módulo')
p(doc, 'Las siguientes secciones describen cada módulo en detalle: componentes, lógica de negocio, reglas de visualización y comportamiento.', italic=True)
sep(doc)

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 0 — ONBOARDING
# ══════════════════════════════════════════════════════════════════════════════
h(doc, 'Módulo 0: Onboarding con Perfil de Inversor', level=1)
p(doc, 'El onboarding es la primera experiencia del usuario dentro del módulo de Inversiones. Se activa una sola vez: la primera vez que accede a Inversiones desde el portafolio general.')
sep(doc)
p(doc, 'Al ingresar, el usuario tiene dos caminos: completar la creación de su cuenta en ese momento, o explorar el producto primero y completarla después. Mientras no finalice el onboarding puede navegar y ver el contenido del módulo (precios, activos, su perfil de inversor), pero no puede operar — ni comprar ni vender — hasta que la cuenta esté completa.')
sep(doc)
p(doc, 'Una vez que completa el onboarding, en sesiones futuras entra directamente al Home (Módulo 1).')
sep(doc)
note(doc, '📌  El perfil de inversor no es solo una formalidad: define qué Packs se le recomiendan, qué badge de volatilidad se destaca y en qué orden aparecen los activos en el Mercado (Módulo 2).')
sep(doc)

h(doc, 'Flujo general', level=2)
p(doc, 'El onboarding tiene 4 etapas secuenciales:')
sep(doc)
mono(doc, """\
PORTAFOLIO GENERAL
       │  Usuario toca card "Inversiones" por primera vez
       ▼
┌──────────────────────────────────────────────────────────┐
│  ETAPA 1 — Propuesta de valor                            │
│  ¿Por qué invertir acá?                                  │
└──────────────────────┬───────────────────────────────────┘
                       ▼
┌──────────────────────────────────────────────────────────┐
│  ETAPA 2 — Cuestionario de perfil de inversor            │
│  5 preguntas · barra de progreso · sin saltar            │
└──────────────────────┬───────────────────────────────────┘
                       ▼
┌──────────────────────────────────────────────────────────┐
│  ETAPA 3 — Resultado del perfil                          │
│  Conservador / Moderado / Dinámico + Pack recomendado    │
└──────────────────────┬───────────────────────────────────┘
                       ▼
┌──────────────────────────────────────────────────────────┐
│  ETAPA 4 — Primera acción                                │
│  "Invertir en el pack" o "Explorar mercado"              │
└──────────────────────┬───────────────────────────────────┘
                       ▼
               HOME / PORTAFOLIO (M1)""")
sep(doc)

h(doc, 'Etapa 1 — Propuesta de valor', level=2)
p(doc, 'Pantalla de bienvenida. Genera confianza y motiva al usuario a completar el onboarding.')
sep(doc)
tbl(doc, ['Elemento', 'Descripción'], [
    ('Ilustración / visual', 'Imagen que transmite crecimiento o seguridad financiera.'),
    ('Título',     '"Hacé crecer tu plata en el mercado global"'),
    ('Subtítulo',  '"Invertí en las empresas más grandes del mundo desde $X. Sin complicaciones."'),
    ('Pill 1', '🌎  Acciones y ETFs de EE.UU.'),
    ('Pill 2', '💵  En pesos o dólares'),
    ('Pill 3', '⚡  Desde cualquier monto'),
    ('CTA principal',   '"Empezar" → avanza a Etapa 2'),
    ('Link secundario', '"¿Cómo funciona?" → modal explicativo (V2)'),
], widths=[1.8, 4.8])
sep(doc)

h(doc, 'Etapa 2 — Cuestionario de perfil de inversor', level=2)
p(doc, '5 preguntas de opción única. Cada respuesta suma puntos que determinan el perfil final. Sin puntajes visibles, sin posibilidad de saltar preguntas. Barra de progreso visible en todo momento.')
sep(doc)
note(doc, '📌  Lenguaje simple y situacional, sin términos financieros técnicos. El objetivo es que cualquier usuario, incluso sin experiencia previa, pueda responder con honestidad.')
sep(doc)

for num, titulo, pregunta, opts in [
    (1, '¿Para qué querés hacer crecer tu plata?', '"Pensá en la razón principal por la que estás acá."', [
        ('A', 'Para que el dólar no me coma los ahorros. Con que le gane a la inflación, me alcanza.', '1'),
        ('B', 'Quiero hacer crecer lo que tengo con el tiempo, sin apuros ni grandes sustos.', '2'),
        ('C', 'Busco que mi plata trabaje duro. Me banco la volatilidad si el premio vale la pena.', '3'),
    ]),
    (2, '¿Alguna vez pusiste tu plata a trabajar?', '"Sé honesto, no hay respuesta incorrecta."', [
        ('A', 'Nunca. El plazo fijo es lo máximo que hice, y tampoco tan seguido.', '1'),
        ('B', 'Algo: dólar, MEP, CEDEARs, cripto... fui probando cosas por mi cuenta.', '2'),
        ('C', 'Sí, invierto con regularidad. Me manejo con los mercados y entiendo los riesgos.', '3'),
    ]),
    (3, '¿Cuándo creés que vas a querer usar esa plata?', '"No importa el monto, pensá en el tiempo."', [
        ('A', 'En menos de un año. Necesito tenerla accesible por las dudas.', '1'),
        ('B', 'En 1 a 3 años. Hay algo concreto en el horizonte, pero no es urgente.', '2'),
        ('C', 'Más de 3 años, o ni lo tengo claro todavía. Esto es para el largo plazo.', '3'),
    ]),
    (4, 'Tu inversión baja un 20% en un mes. ¿Qué hacés?', '"Imaginátelo de verdad, no respondas lo que creés que deberías decir."', [
        ('A', 'Salgo. Prefiero cortar la pérdida antes de que se ponga peor.', '1'),
        ('B', 'Me quedo quieto. Sé que los mercados suben y bajan, espero que se recupere.', '2'),
        ('C', 'Compro más. Una caída así para mí es una oportunidad, no una alarma.', '3'),
    ]),
    (5, '¿Qué parte de tus ahorros estás pensando invertir?', '"Del total que tenés ahorrado hoy, ¿con cuánto querés arrancar?"', [
        ('A', 'Poco, menos del 20%. Quiero probar cómo me va antes de meterme más.', '1'),
        ('B', 'Una parte razonable, entre 20% y 50%. Tengo algo de idea de lo que hago.', '2'),
        ('C', 'La mayor parte, más del 50%. Le tengo confianza a lo que hago.', '3'),
    ]),
]:
    p(doc, f'Pregunta {num} — {titulo}', bold=True)
    p(doc, pregunta)
    tbl(doc, ['Opción', 'Respuesta', 'Puntaje'], opts, widths=[0.5, 5.2, 0.9])
    sep(doc)

h(doc, 'Lógica de scoring', level=3)
p(doc, 'Cada pregunta suma entre 1 y 3 puntos. Puntaje total: 5 a 15.')
tbl(doc, ['Puntaje total', 'Perfil', 'Descripción'], [
    ('5 — 8',  'Conservador', 'Baja tolerancia al riesgo. Prioridad: preservar capital.'),
    ('9 — 12', 'Moderado',    'Riesgo balanceado. Busca crecimiento con estabilidad.'),
    ('13 — 15','Dinámico',    'Alta tolerancia al riesgo. Foco en crecimiento agresivo.'),
], widths=[1.2, 1.2, 4.2])
sep(doc)
note(doc, '📌  En empate entre rangos → se asigna el perfil más conservador. El usuario puede cambiar su perfil desde Configuración en cualquier momento.')
sep(doc)

h(doc, 'Etapa 3 — Resultado del perfil', level=2)
p(doc, 'Pantalla que le muestra al usuario su perfil, qué significa, y el Pack recomendado como primer destino de inversión.')
sep(doc)
for perfil, icono, badge, pack, activos in [
    ('CONSERVADOR', 'Azul · escudo',   '🔵  Baja', 'Pack Estabilidad (ETFs de bonos + empresas consolidadas)',          'SPY, VTI, JNJ, KO, PG'),
    ('MODERADO',    'Violeta · balanza','🟣  Media', 'Pack Crecimiento Balanceado (ETFs globales + sectores líderes)',   'QQQ, MSFT, AAPL, V, UNH'),
    ('DINÁMICO',    'Naranja · cohete', '🟠  Alta',  'Pack Alto Potencial (tech + crecimiento + ETFs temáticos)',        'NVDA, TSLA, META, ARKK, AMD'),
]:
    p(doc, f'Perfil {perfil}', bold=True)
    tbl(doc, ['Campo', 'Valor'], [
        ('Ícono / color',        icono),
        ('Título',               f'"Sos un inversor {perfil.capitalize()}"'),
        ('Badge de volatilidad', badge),
        ('Pack recomendado',     pack),
        ('Activos sugeridos',    activos),
    ], widths=[2.0, 4.6])
    sep(doc)

p(doc, 'Elementos comunes a los 3 resultados:', bold=True)
tbl(doc, ['Elemento', 'Descripción'], [
    ('Card del Pack',      'Nombre · rendimiento 12M · cantidad de activos · badge de volatilidad. Igual que en M2.'),
    ('CTA principal',      '"Invertir en este Pack" → detalle del Pack (M2) con flujo de compra (M4) pre-cargado.'),
    ('CTA secundario',     '"Explorar todo el mercado" → Módulo 2 sin filtro.'),
    ('Link de perfil',     '"Ver mi perfil completo" → modal con los 3 perfiles y opción de cambiar.'),
    ('Cambiar respuestas', 'Vuelve al inicio del cuestionario.'),
], widths=[2.0, 4.6])
sep(doc)

h(doc, 'Etapa 4 — Primera acción', level=2)
p(doc, 'El usuario llega al Home (M1) con el portafolio vacío. Si tocó "Invertir en el Pack recomendado" en Etapa 3, llega directamente al detalle de ese Pack dentro del Módulo 2.')
note(doc, '📌  No hay pantalla de "felicitaciones" ni confeti separado. El resultado del perfil (Etapa 3) funciona como cierre emotivo. La primera acción real del usuario define el inicio de la relación con el producto.')
sep(doc)

h(doc, 'Reglas generales', level=2)
tbl(doc, ['Regla', 'Descripción'], [
    ('Unicidad',            'El cuestionario se muestra una sola vez. Una vez completado, el perfil queda guardado.'),
    ('Reasignación manual', 'El usuario puede cambiar su perfil desde Configuración → Perfil de inversor en cualquier momento.'),
    ('Sin omisión',         'No hay opción de saltear el cuestionario. El CTA de Etapa 1 solo avanza al cuestionario.'),
    ('Persistencia',        'Si el usuario cierra la app durante el cuestionario, al volver retoma desde donde lo dejó.'),
    ('Modo exploración',    'Antes de completar el onboarding el usuario puede navegar y ver contenido, pero no puede operar.'),
    ('Sin cuenta nueva',    'El onboarding no crea una cuenta nueva. Solo activa el módulo de Inversiones y define el perfil.'),
], widths=[2.0, 4.6])
sep(doc)

h(doc, 'Cómo impacta el perfil en el resto del producto', level=2)
tbl(doc, ['Módulo', 'Impacto'], [
    ('M2 — Mercado / Packs', 'El Pack recomendado aparece primero. Badge de volatilidad destacado con el color del perfil.'),
    ('M2 — Acciones',        'Badge "Recomendado para vos" en activos compatibles con el perfil del usuario.'),
    ('M3 — Detalle activo',  '"Compatible con tu perfil Moderado" (V2).'),
    ('M1 — Home vacío',      'CTA "Explorar mercado" lleva al Mercado con filtro de recomendados pre-aplicado.'),
], widths=[2.0, 4.6])
sep(doc)

h(doc, 'V1 vs V2 — Onboarding', level=2)
tbl(doc, ['Funcionalidad', 'V1', 'V2'], [
    ('Propuesta de valor (Etapa 1)',            '✅', ''),
    ('Cuestionario 5 preguntas (Etapa 2)',      '✅', ''),
    ('Resultado + Pack recomendado (Etapa 3)',  '✅', ''),
    ('CTA de primera acción (Etapa 4)',         '✅', ''),
    ('Persistencia del progreso',               '✅', ''),
    ('Modo exploración sin operar',             '✅', ''),
    ('Reasignación desde Configuración',        '✅', ''),
    ('Badge "Recomendado" en Mercado',          '✅', ''),
    ('Modal "¿Cómo funciona?" (Etapa 1)',       '', '🔜'),
    ('Re-evaluación periódica del perfil',      '', '🔜'),
    ('5 perfiles en lugar de 3',                '', '🔜'),
    ('Compatibilidad de perfil en M3',          '', '🔜'),
    ('Cuestionario con preguntas regulatorias', '', '🔜'),
], widths=[4.0, 0.8, 0.8])
sep(doc)

h(doc, 'Pantallas del flujo (resumen)', level=2)
p(doc, 'Total: 8 pantallas para completar el onboarding.')
sep(doc)
mono(doc, """\
 Pantalla 1   Propuesta de valor
              Ilustración + título + 3 pills + CTA "Empezar"
              ──────────────────────────────────────────────
 Pantalla 2   Pregunta 1/5 — ¿Para qué querés hacer crecer tu plata?  (20%)
              ──────────────────────────────────────────────
 Pantalla 3   Pregunta 2/5 — ¿Alguna vez pusiste tu plata a trabajar?  (40%)
              ──────────────────────────────────────────────
 Pantalla 4   Pregunta 3/5 — ¿Cuándo creés que vas a querer usar esa plata?  (60%)
              ──────────────────────────────────────────────
 Pantalla 5   Pregunta 4/5 — Tu inversión baja 20% en un mes. ¿Qué hacés?  (80%)
              ──────────────────────────────────────────────
 Pantalla 6   Pregunta 5/5 — ¿Qué parte de tus ahorros pensás invertir?  (100%)
              ──────────────────────────────────────────────
 Pantalla 7   Resultado del perfil
              Conservador / Moderado / Dinámico + Pack recomendado + 2 CTAs
              ──────────────────────────────────────────────
 Pantalla 8   Home / Portafolio (M1) — estado vacío
              Primera vez que el usuario ve el producto real""")
sep(doc)

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULOS 1–7 (contenido original)
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(doc, 'Módulo 1: Home / Portafolio', level=1)
p(doc, 'Pantalla central del módulo. El usuario llega acá desde el Portfolio general tocando el card de Inversiones.')
sep(doc)

h(doc, 'Header — Balance y rendimiento', level=2)
tbl(doc, ['Campo', 'Cálculo'], [
    ('Balance total',    'Suma de (precio actual × cantidad) + dividendos de todas las posiciones abiertas, en USD'),
    ('Variación en USD', 'Valor actual total − total invertido (precio promedio ponderado × cantidad)'),
    ('Variación en %',   '(variación USD / total invertido) × 100'),
], widths=[2.2, 4.4])
sep(doc)
p(doc, 'Períodos disponibles:', bold=True)
tbl(doc, ['Período', 'Cálculo'], [
    ('Diario',    'Cierre D-1 vs precio en tiempo real'),
    ('Semana',    'Últimos 7 días corridos'),
    ('Mes',       'Últimos 30 días corridos'),
    ('Año',       '1° enero del año en curso hasta hoy'),
    ('Histórico', 'Desde la primera compra del usuario'),
], widths=[1.5, 5.1])
sep(doc)
p(doc, 'El label del período y el P&L se actualizan dinámicamente al cambiar el selector. Verde si positivo, rojo si negativo.')
sep(doc)

h(doc, 'Barra de distribución', level=2)
p(doc, 'Barra segmentada proporcional al valor actual de cada posición. Orden: mayor a menor valor (izquierda → derecha). Cada segmento tiene el color fijo del activo. Leyenda con ticker + % debajo. Máximo 5 segmentos; el resto se agrupa.')
sep(doc)
tbl(doc, ['Vista', 'Comportamiento'], [
    ('Todos',       'Lista plana por valor descendente (default)'),
    ('Sectores',    'Agrupado por sector (Tecnología, Automotriz, etc.)'),
    ('Tipo',        'Agrupado por tipo de activo (Acciones / ETFs)'),
    ('Dividendos',  'Filtra y muestra solo activos que pagan dividendo'),
], widths=[1.5, 5.1])
sep(doc)

h(doc, 'Lista de posiciones', level=2)
p(doc, 'Ordenada por valor actual descendente. Máximo 5 posiciones visibles + botón "Ver más".')
p(doc, 'Cada ítem: logo · nombre completo · cantidad de acciones · valor actual USD · rendimiento % del período seleccionado.')
sep(doc)
p(doc, 'Sin posiciones abiertas: ilustración + texto + botón "Explorar mercado" como único CTA.')
sep(doc)

h(doc, 'Indicador contextual de orden activa', level=2)
p(doc, 'Banner que aparece solo cuando hay una orden en curso. Si no hay órdenes activas, no se muestra. Al tocar → abre el Centro de Órdenes filtrado por esa orden.')
sep(doc)
tbl(doc, ['Estado', 'Texto', 'Color'], [
    ('Procesando', '"Compra en proceso · $500 ARS"',   '🟡 Amarillo'),
    ('Completada', '"Compra completada · $500 ARS"',   '🟢 Verde'),
    ('Cancelada',  '"Compra cancelada · $500 ARS"',    '🔴 Rojo'),
], widths=[1.5, 3.5, 1.6])
sep(doc)
note(doc, '📌  El estado "Completada" desaparece automáticamente después de unos segundos. Solo se muestra si hay al menos una orden en estado Procesando.')
sep(doc)

h(doc, 'Vista Gráfico ✅ V1', level=2)
p(doc, 'Permite ver el rendimiento del portafolio en el tiempo y comparar curvas. Resuelve la pregunta: "¿Cómo le está yendo a cada acción, y cómo me está yendo vs el mercado?"')
sep(doc)
tbl(doc, ['Estado', 'Comportamiento'], [
    ('Sin selección (default)', 'Curva del portafolio total. Arranca en 0% y termina en el rendimiento del período. Línea punteada en 0% de referencia.'),
    ('1 chip seleccionado',     'Desaparece la curva del portafolio. Aparece la curva del activo con su color propio.'),
    ('Múltiples chips',         'Múltiples curvas simultáneas. El eje Y se recalcula dinámicamente. Límite: 4 selecciones.'),
    ('Límite alcanzado',        'Los chips no seleccionados se ponen a 30% de opacidad y dejan de ser tapeables.'),
    ('Scrubber interactivo',    'Al arrastrar el dedo: línea vertical + punto en cada curva + % en el header del card.'),
], widths=[2.2, 4.4])
sep(doc)
p(doc, 'Chips "Mis acciones": todas las posiciones del usuario con su rendimiento en el período.')
p(doc, 'Chips "Comparar con": benchmarks externos fijos (S&P 500, etc.). Consumen del mismo límite de 4.')
sep(doc)

# M2
h(doc, 'Módulo 2: Mercado', level=1)
p(doc, 'Pantalla de descubrimiento. El usuario explora activos disponibles para invertir.')
p(doc, 'Acceso: botón "Explorar mercado" desde el Home vacío · botón de compra desde el Home · ícono de búsqueda en el header.')
sep(doc)

h(doc, 'Buscador', level=2)
p(doc, 'Se activa desde el ícono de lupa en el header. Búsqueda en tiempo real sobre todos los activos (Packs + ETFs + Acciones). Busca por nombre y por ticker.')
sep(doc)

h(doc, 'Packs', level=2)
p(doc, 'Carteras curadas. El usuario invierte en varias acciones a la vez para diversificar según su perfil de inversor.')
sep(doc)
tbl(doc, ['Campo', 'Descripción'], [
    ('Nombre del pack',   'Nombre y descripción breve (1 línea)'),
    ('Rendimiento',       'Últimos 12 meses en %'),
    ('Composición',       'Avatares apilados de los activos + cantidad ("4 activos")'),
    ('Badge volatilidad', 'Baja / Media / Alta'),
    ('CTA',               '"Ver detalle" → abre detalle del pack'),
], widths=[2.0, 4.6])
sep(doc)

h(doc, 'Índices (ETFs)', level=2)
p(doc, 'Fondos cotizados que agrupan múltiples activos.')
tbl(doc, ['Campo', 'Descripción'], [
    ('Logo + nombre + ticker', 'Identificación del ETF'),
    ('Descripción',             '"Las 500 empresas más grandes de EE.UU."'),
    ('Variación del día',       '% respecto al cierre D-1'),
], widths=[2.0, 4.6])
sep(doc)

h(doc, 'Acciones', level=2)
p(doc, 'Filtros disponibles: Todos · Favoritos · Populares · Tecnología · Con dividendos')
tbl(doc, ['Campo', 'Descripción'], [
    ('Logo + nombre completo', 'Identificación'),
    ('Ticker · Sector',         'Categoría del activo'),
    ('Variación del día',       '% respecto al cierre D-1'),
], widths=[2.0, 4.6])
sep(doc)

# M3
h(doc, 'Módulo 3: Detalle del activo', level=1)
p(doc, 'Pantalla individual de un activo. Accesible desde el portafolio (si tiene posición) y desde el Mercado. Incluye ícono para agregar a favoritos.')
sep(doc)

h(doc, 'Header', level=2)
tbl(doc, ['Campo', 'Descripción'], [
    ('Logo + nombre + ticker', 'Identificación del activo'),
    ('Sector',                  'Categoría del activo'),
    ('Precio actual',           'En USD, con 2 decimales'),
    ('Variación del día',       '% respecto al cierre D-1, con flecha si subió o bajó'),
], widths=[2.0, 4.6])
sep(doc)

h(doc, 'Gráfico de precio ✅ V1', level=2)
p(doc, 'Muestra el precio histórico del activo. Por defecto en 1D. Scrubbing táctil: al deslizar muestra precio exacto + fecha/hora.')
sep(doc)
tbl(doc, ['Período', 'Eje X'], [
    ('1D',       'Cada 30 minutos'),
    ('1W',       'Lun · Mar · Mié · Jue · Vie'),
    ('1M',       'S1 · S2 · S3 · S4'),
    ('1Y',       'Todos los meses'),
    ('Histórico','Desde el inicio'),
], widths=[1.5, 5.1])
sep(doc)
note(doc, '📌  V2: se agrega un marcador visual del precio promedio de compra del usuario sobre la curva.')
sep(doc)

h(doc, 'Descripción y Noticias ✅ V1', level=2)
p(doc, 'Descripción: breve descripción de la empresa con tags relacionados.')
p(doc, 'Noticias: feed de noticias vinculadas al activo o al sector. Cada ítem: título · fuente · tiempo transcurrido · tag temático.')
note(doc, '📌  V1: noticias provistas por Alpaca, traducidas al español. V2: noticias financieras en tiempo real con proveedor externo.')
sep(doc)

h(doc, 'Valoración ⏳ Post-lanzamiento (requiere proveedor externo)', level=2)
p(doc, 'Datos financieros fundamentales. No disponible en V1.')
sep(doc)
tbl(doc, ['Campo técnico', 'Nombre en UI', 'Tooltip'], [
    ('Market Cap',        'Tamaño de la empresa',      '"Es lo que vale la empresa entera si alguien quisiera comprarla hoy."'),
    ('Ratio P/E',         'Puntaje de valoración',     '"Ayuda a saber si la acción está cara o barata."'),
    ('Dividendo anual',   'Pago al accionista',        '"El % de dinero extra que la empresa te paga cada año por tener sus acciones."'),
    ('Rango 52 semanas',  'Rango del último año',      '"El precio más bajo y el más alto de los últimos 12 meses."'),
    ('Margen neto',       'Ganancia limpia',           '"El % que le queda al negocio después de pagar todos sus gastos e impuestos."'),
], widths=[1.8, 1.8, 3.0])
sep(doc)

h(doc, 'Posición del usuario ✅ V1', level=2)
p(doc, 'Visible solo si el usuario tiene o tuvo posición abierta en este activo.')
tbl(doc, ['Campo', 'Cálculo'], [
    ('Valor actual',              'Precio actual × cantidad total, en USD'),
    ('Variación USD',             'Valor actual − total invertido'),
    ('Variación %',               '(variación USD / total invertido) × 100'),
    ('Cantidad de acciones',      'Total compras − total ventas'),
    ('Precio promedio de compra', 'Promedio ponderado de todas las órdenes ejecutadas'),
    ('Total invertido',           'Suma de (precio × cantidad) de cada orden de compra'),
], widths=[2.4, 4.2])
sep(doc)

h(doc, 'Historial de transacciones ✅ V1', level=2)
p(doc, 'Transacciones del usuario para ese ticker: tipo · estado · fecha · cantidad · precio de operación. Al tocar un ítem se expande el detalle completo.')
sep(doc)

h(doc, 'Botones de acción (fijo abajo)', level=2)
tbl(doc, ['Estado', 'Botones'], [
    ('No tiene el activo', 'Comprar'),
    ('Tiene el activo',    'Comprar más + Vender'),
], widths=[2.0, 4.6])
sep(doc)

# M4
h(doc, 'Módulo 4: Flujo de compra', level=1)
note(doc, '📌  Tipo de orden en V1: Orden de mercado. Se ejecuta al precio disponible en el momento de la confirmación. No hay precio garantizado.')
sep(doc)
tbl(doc, ['Paso', 'Descripción'], [
    ('1', 'El usuario selecciona la moneda en la que quiere comprar: ARS o USD'),
    ('2', 'Ve cuánto representa 1 acción en la moneda seleccionada'),
    ('3', 'Ve el saldo disponible según el método de pago seleccionado'),
    ('4', 'Ingresa el monto. Si supera el saldo disponible → error "Saldo insuficiente"'),
    ('5', 'Revisa el resumen: activo + logo · cantidad estimada de acciones · monto (comisión incluida) · % de comisión'),
    ('6', 'Confirmar compra → dispara la orden → va al Módulo 6'),
], widths=[0.4, 6.2])
sep(doc)

# M5
h(doc, 'Módulo 5: Flujo de venta', level=1)
note(doc, '📌  Tipo de orden en V1: Orden de mercado. Se ejecuta al precio disponible en el momento de la confirmación. No hay precio garantizado.')
sep(doc)
tbl(doc, ['Paso', 'Descripción'], [
    ('1', 'El usuario selecciona la moneda en la que quiere recibir los fondos: ARS o USD'),
    ('2', 'Ve cuánto representa 1 acción en la moneda seleccionada'),
    ('3', 'Ve el saldo disponible en el activo'),
    ('4', 'Ingresa el monto. Si supera el saldo disponible → error "Saldo insuficiente"'),
    ('5', 'Revisa el resumen: activo + logo · cantidad estimada a vender · monto (comisión incluida) · % de comisión'),
    ('6', 'Confirmar venta → dispara la orden → va al Módulo 6'),
], widths=[0.4, 6.2])
sep(doc)

# M6
h(doc, 'Módulo 6: Centro de órdenes y actividad', level=1)
p(doc, 'Registro centralizado de todas las órdenes de inversión del usuario. Las órdenes pasan por estados. Este módulo permite saber en todo momento qué pasó con cada una.')
p(doc, 'Acceso: ícono de actividad en el header de Inversiones Home.')
sep(doc)

h(doc, 'V1 — Órdenes de mercado', level=2)
mono(doc, """\
ENVIADA
   └→ PROCESANDO  (en ejecución en el mercado)
         ├→ COMPLETA   (ejecutada exitosamente)
         └→ CANCELADA  (no se pudo ejecutar)""")
sep(doc)
tbl(doc, ['Estado', 'Color', 'Qué ve el usuario'], [
    ('Procesando', '🟡 Amarillo', '"Tu orden está siendo procesada"'),
    ('Completa',   '🟢 Verde',    'Precio de ejecución real + acciones recibidas/vendidas'),
    ('Cancelada',  '🔴 Rojo',     'Motivo + botón "Reintentar"'),
], widths=[1.4, 1.4, 3.8])
sep(doc)

h(doc, 'Lista de órdenes', level=2)
p(doc, 'Orden: cronológico descendente (más reciente primero).')
p(doc, 'Cada ítem: tipo de movimiento (Compra/Venta) · logo + ticker · fecha y hora · monto USD · cantidad de acciones · estado con color.')
sep(doc)

h(doc, 'Vista detalle de una orden', level=2)
tbl(doc, ['Campo', 'Descripción'], [
    ('Tipo de orden',           'Mercado (V1)'),
    ('Activo',                  'Nombre + ticker'),
    ('Fecha y hora de creación','Timestamp exacto'),
    ('Fecha y hora de ejecución','Timestamp exacto (si aplica)'),
    ('Precio de referencia',    'Al momento de confirmar'),
    ('Precio de ejecución real','Precio al que se ejecutó (puede diferir)'),
    ('Cantidad de acciones',    'Definitiva post-ejecución'),
    ('Monto final',             'Monto y moneda ingresados por el usuario'),
    ('Comisión',                '% aplicado'),
    ('Estado + motivo',         'Si fue cancelada, el motivo'),
], widths=[2.4, 4.2])
sep(doc)

h(doc, 'V2 — Órdenes límite', level=2)
mono(doc, """\
CREADA
   └→ ACTIVA      (esperando que el activo alcance el precio objetivo)
         ├→ EJECUTADA  (se alcanzó el precio)
         ├→ VENCIDA    (el período venció sin alcanzar el precio)
         └→ CANCELADA  (el usuario la canceló manualmente)""")
sep(doc)
tbl(doc, ['Estado', 'Color', 'Qué ve el usuario'], [
    ('Activa',    '🟡 Amarillo', 'Precio objetivo + distancia % al precio actual + tiempo restante'),
    ('Ejecutada', '🟢 Verde',    'Precio de ejecución + acciones recibidas'),
    ('Vencida',   '⚪ Gris',     '"La orden venció" + opción de recrear'),
    ('Cancelada', '🔴 Rojo',     '"Cancelada por vos"'),
], widths=[1.4, 1.4, 3.8])
sep(doc)

h(doc, 'V2 — Órdenes after market', level=2)
p(doc, 'El usuario puede crear órdenes de compra fuera del horario regular (antes de las 9:30am o después de las 4:00pm ET). La orden se ejecuta durante la ventana after market (4pm–8pm ET).')
sep(doc)
tbl(doc, ['Estado', 'Qué significa'], [
    ('Programada',    'La orden fue creada fuera del horario de mercado y queda en cola'),
    ('En ventana',    'Está en el período de after market, activamente buscando ejecución'),
    ('Ejecutada',     'Se ejecutó durante la ventana after market'),
    ('No ejecutada',  'La ventana cerró sin encontrar contraparte'),
], widths=[1.8, 4.8])
sep(doc)

# M7
h(doc, 'Módulo 7: Resultados (Posiciones cerradas)', level=1)
p(doc, 'Historial de posiciones liquidadas — activos que el usuario compró y ya vendió en su totalidad. Muestra el resultado final de cada operación cerrada.')
p(doc, 'Acceso: ícono en el header de Inversiones Home.')
sep(doc)

h(doc, 'Header', level=2)
tbl(doc, ['Campo', 'Cálculo'], [
    ('G/P total realizado', 'Suma de todas las ganancias y pérdidas de posiciones cerradas, en USD'),
    ('Rendimiento total %',  'G/P total / suma de lo invertido en esas posiciones × 100'),
], widths=[2.2, 4.4])
sep(doc)

h(doc, 'Lista de posiciones cerradas', level=2)
p(doc, 'Orden: cronológico descendente por fecha de venta (la más reciente primero).')
sep(doc)
tbl(doc, ['Vista', 'Campos'], [
    ('Colapsada', 'Logo + ticker + nombre · fecha de venta · G/P en USD (verde/rojo) · % de rendimiento · inversión inicial · monto de venta'),
    ('Expandida (al tocar)', 'Unidades · días promedio de tenencia · P&L por unidad · precio promedio de compra · precio de venta por unidad'),
], widths=[1.8, 4.8])
sep(doc)
p(doc, 'Filtros: por fecha · por activo.')
sep(doc)

# ══════════════════════════════════════════════════════════════════════════════
# MÓDULO 8 — REPORTES Y COMUNICACIONES
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h(doc, 'Módulo 8: Reportes y Comunicaciones', level=1)
p(doc, 'Este módulo cubre dos features relacionados pero distintos: la descarga de reportes del historial de inversiones, y el envío de comunicaciones periódicas con el estado de la posición del usuario. Ambos tienen como objetivo que el usuario mantenga visibilidad sobre su actividad sin necesidad de abrir la app.')
sep(doc)

# ─── PARTE A — REPORTES ────────────────────────────────────────────────────
h(doc, 'Parte A — Reportes descargables', level=2)
p(doc, 'El usuario puede descargar un PDF con el detalle de su actividad inversora para un período dado. Hay dos tipos de reporte: mensual y anual.')
sep(doc)
note(doc, '📌  El reporte no es solo un historial de transacciones — incluye P&L realizado y no realizado, dividendos cobrados y comisiones pagadas. Es el documento que el usuario necesita para su declaración de impuestos o para entender qué hizo con su plata en un período.')
sep(doc)

h(doc, 'Punto de acceso', level=3)
tbl(doc, ['Desde dónde', 'Cómo'], [
    ('Módulo 7 — Resultados',      'Botón "Descargar reporte" en el header de la pantalla de posiciones cerradas.'),
    ('Módulo 6 — Centro de órdenes','Botón "Descargar reporte" en el header del historial de órdenes.'),
    ('Configuración / Perfil',     'Sección "Mis reportes" con acceso a todos los reportes generados.'),
], widths=[2.2, 4.4])
sep(doc)

h(doc, 'Tipos de reporte', level=3)
p(doc, 'Reporte mensual', bold=True)
tbl(doc, ['Campo', 'Detalle'], [
    ('Disponibilidad', 'A partir del día 1 del mes siguiente al período cerrado.'),
    ('Período',        'Del 1° al último día del mes calendario.'),
    ('Generación',     'Automática al cierre de cada mes. El usuario solo descarga.'),
    ('Formato',        'PDF. Nombre: "Reporte_[mes]_[año].pdf"'),
], widths=[2.2, 4.4])
sep(doc)
p(doc, 'Reporte anual', bold=True)
tbl(doc, ['Campo', 'Detalle'], [
    ('Disponibilidad', 'A partir del 1° de enero del año siguiente al período cerrado.'),
    ('Período',        'Del 1° de enero al 31 de diciembre.'),
    ('Generación',     'Automática al cierre del año. El usuario solo descarga.'),
    ('Formato',        'PDF. Nombre: "Reporte_anual_[año].pdf"'),
], widths=[2.2, 4.4])
sep(doc)

h(doc, 'Contenido del reporte', level=3)
p(doc, 'Ambos tipos comparten la misma estructura. El anual agrega la sección de resumen general del año.')
sep(doc)
tbl(doc, ['Sección', 'Contenido', 'Mensual', 'Anual'], [
    ('Encabezado',
     'Nombre del usuario, período, fecha de generación, moneda base (USD).',
     '✅', '✅'),
    ('Resumen del período',
     'Balance inicial · Balance final · Variación total en USD y % · Dividendos cobrados · Comisiones pagadas.',
     '✅', '✅'),
    ('Posiciones abiertas al cierre',
     'Por activo: cantidad · precio promedio de compra · precio al cierre del período · valor actual · P&L no realizado en USD y %.',
     '✅', '✅'),
    ('Transacciones del período',
     'Tabla cronológica: fecha · tipo · activo · cantidad · precio de ejecución · monto total · comisión.',
     '✅', '✅'),
    ('Posiciones cerradas en el período',
     'Activos comprados y vendidos en su totalidad: precio de entrada · precio de salida · G/P realizado en USD y %.',
     '✅', '✅'),
    ('Dividendos cobrados',
     'Por activo: fecha de cobro · monto en USD.',
     '✅', '✅'),
    ('Resumen anual',
     'P&L realizado total del año · P&L no realizado al 31/12 · rendimiento anual del portafolio vs benchmark (S&P 500).',
     '—', '✅'),
], widths=[1.8, 3.6, 0.7, 0.7])
sep(doc)
note(doc, '📌  Los reportes se generan del lado del servidor. Si no hubo actividad en un mes, el reporte se genera igual (con posiciones abiertas y balance) pero sin sección de transacciones.')
sep(doc)

h(doc, 'Pantalla "Mis reportes"', level=3)
p(doc, 'Accesible desde Configuración / Perfil. Lista todos los reportes generados, ordenados cronológicamente descendente.')
tbl(doc, ['Elemento', 'Descripción'], [
    ('Ítem de la lista', 'Tipo (Mensual / Anual) · período · fecha de generación · botón "Descargar".'),
    ('Estado',           '"Disponible" si ya fue generado · "En preparación" si el mes acaba de cerrar.'),
    ('Descarga',         'Toca "Descargar" → abre el PDF en el visor nativo del dispositivo.'),
    ('Primer reporte',   'Se genera al cierre del primer mes en que el usuario hizo su primera transacción.'),
], widths=[2.2, 4.4])
sep(doc)

h(doc, 'V1 vs V2 — Reportes', level=3)
tbl(doc, ['Funcionalidad', 'V1', 'V2'], [
    ('Reporte mensual descargable en PDF',                    '✅', ''),
    ('Reporte anual descargable en PDF',                      '✅', ''),
    ('Pantalla "Mis reportes" en Configuración',              '✅', ''),
    ('Posiciones abiertas al cierre en el reporte',           '✅', ''),
    ('Transacciones del período en el reporte',               '✅', ''),
    ('Posiciones cerradas + G/P en el reporte',               '✅', ''),
    ('Dividendos cobrados en el reporte',                     '✅', ''),
    ('Resumen anual con benchmark (S&P 500)',                 '',   '🔜'),
    ('Descarga en formato Excel / CSV',                       '',   '🔜'),
    ('Reporte con formato para declaración AFIP',             '',   '🔜'),
    ('Notificación push cuando el reporte del mes está listo','',   '🔜'),
], widths=[4.4, 0.7, 0.7])
sep(doc)

# ─── PARTE B — COMUNICACIONES ──────────────────────────────────────────────
h(doc, 'Parte B — Comunicaciones periódicas', level=2)
p(doc, 'El sistema envía comunicaciones automáticas al usuario con el estado de su portafolio: cuánto tiene, cuánto ganó o perdió (realizado y no realizado), y cómo evolucionó. Son mensajes proactivos — el usuario no tiene que abrir la app para enterarse cómo le está yendo.')
sep(doc)
note(doc, '📌  El objetivo es mantener al usuario conectado con su inversión sin que tenga que acordarse de revisar la app. Una buena comunicación periódica genera confianza, reduce la ansiedad y aumenta la retención.')
sep(doc)

h(doc, 'Canal de entrega', level=3)
tbl(doc, ['Canal', 'Descripción', 'V1', 'V2'], [
    ('Push notification', 'Notificación en el dispositivo. Al tocar, abre directamente el Home (M1).', '✅', ''),
    ('Email',             'Mail con el resumen completo formateado. Incluye tabla de posiciones y CTA.',  '✅', ''),
    ('In-app banner',     'Banner en el Home (M1) la primera vez que abre la app después del envío.',    '',   '🔜'),
], widths=[1.8, 3.8, 0.6, 0.6])
sep(doc)

h(doc, 'Frecuencia', level=3)
tbl(doc, ['Frecuencia', 'Descripción', 'V1', 'V2'], [
    ('Mensual (fija)',  'Se envía el primer día hábil de cada mes con el resumen del mes anterior. No configurable en V1.',          '✅', ''),
    ('Semanal',        'Resumen los lunes con la variación de los últimos 7 días.',                                                   '', '🔜'),
    ('Trimestral',     'Resumen cada 3 meses con una vista más amplia del período.',                                                  '', '🔜'),
    ('Configurable',   'El usuario elige la frecuencia desde Configuración → Notificaciones.',                                        '', '🔜'),
], widths=[1.8, 3.8, 0.6, 0.6])
sep(doc)
note(doc, '📌  En V1 la frecuencia es mensual fija. No se envía si el usuario no tuvo actividad ni posiciones abiertas en el período.')
sep(doc)

h(doc, 'Contenido de la comunicación', level=3)
p(doc, 'Estructura fija, igual en push y en email (adaptada al formato de cada canal):')
sep(doc)
mono(doc, """\
ASUNTO (email) / TÍTULO (push):
  "Tu resumen de inversiones — [Mes] [Año]"
  Ejemplo: "Tu resumen de inversiones — Febrero 2025"

────────────────────────────────────────────────────────────
SECCIÓN 1 — Balance actual
  Valor total del portafolio en USD al cierre del período
  Variación vs. el mes anterior en USD y %
  (verde si subió, rojo si bajó)

SECCIÓN 2 — Ganancia no realizada
  P&L de las posiciones que todavía tenés abiertas
  = valor actual total − total invertido
  Detalle por activo: ticker · valor actual · P&L USD · P&L %

SECCIÓN 3 — Ganancia realizada en el período
  G/P de las posiciones que cerraste en el mes
  Si no cerró ninguna: "No vendiste posiciones este mes."

SECCIÓN 4 — Actividad del mes
  Cantidad de órdenes ejecutadas (compras + ventas)
  Dividendos cobrados en el período (si aplica)

CTA principal:   "Ver mi portafolio"   → abre Home (M1)
CTA secundario (solo email): "Descargar reporte del mes" → PDF mensual
────────────────────────────────────────────────────────────""")
sep(doc)

h(doc, 'Reglas de envío', level=3)
tbl(doc, ['Regla', 'Descripción'], [
    ('Sin actividad',  'Si el usuario no tiene posiciones abiertas ni cerró nada en el mes, no se envía la comunicación.'),
    ('Primer envío',   'Se genera al cierre del primer mes en que el usuario tuvo al menos una transacción.'),
    ('Opt-out',        'El usuario puede desactivar desde Configuración → Notificaciones. En V1: todo o nada. En V2: configurable por canal y frecuencia.'),
    ('Zona horaria',   'Envío procesado en UTC-3 (hora argentina). Email a las 9:00am del primer día hábil del mes.'),
    ('Moneda',         'Siempre en USD. Si el usuario operó en ARS, se convierte al tipo de cambio del día de la transacción.'),
], widths=[2.0, 4.6])
sep(doc)

h(doc, 'V1 vs V2 — Comunicaciones', level=3)
tbl(doc, ['Funcionalidad', 'V1', 'V2'], [
    ('Comunicación mensual automática (push + email)',         '✅', ''),
    ('Balance actual + variación vs. mes anterior',           '✅', ''),
    ('Ganancia no realizada por posición abierta',            '✅', ''),
    ('Ganancia realizada en el período',                      '✅', ''),
    ('Actividad del mes (órdenes + dividendos)',              '✅', ''),
    ('CTA "Descargar reporte del mes" en el email',           '✅', ''),
    ('Opt-out desde Configuración',                           '✅', ''),
    ('Comunicación semanal o trimestral',                     '',   '🔜'),
    ('Frecuencia configurable por el usuario',                '',   '🔜'),
    ('In-app banner en Home al primer acceso post-envío',     '',   '🔜'),
    ('Comparación vs. benchmark (S&P 500) en el resumen',    '',   '🔜'),
    ('Configuración por canal (push vs. email por separado)', '',   '🔜'),
    ('Rendimiento acumulado desde el inicio en el resumen',  '',   '🔜'),
], widths=[4.4, 0.7, 0.7])
sep(doc)

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f'✅  Guardado: {OUT}')
print(f'    Párrafos totales: {len([p for p in doc.paragraphs if p.text.strip()])}')
